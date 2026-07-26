"""Gera os documentos editáveis da Prioridade 1 em DOCX.

Os PDFs são produzidos depois pelo fluxo de renderização/validação do projeto.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BRANDING = ROOT / "branding"
PROPOSAL = ROOT / "proposta-comercial"

CREATED = "22 de julho de 2026"
APP_URL = "https://aluguel-casas-anderton.netlify.app"
EMAIL = "andertonaluguel@gmail.com"

COVER = "14322A"
COVER_LIGHT = "1F4339"
PAPER = "F4F6F3"
CARD = "FFFFFF"
INK = "1C2620"
INK_SOFT = "5C6B63"
BRASS = "B8863C"
BRASS_DEEP = "8C631F"
BRASS_SOFT = "F1E4C8"
RUST = "A23B2E"
RUST_SOFT = "F4DCD7"
LINE = "DDE4DE"
BLUE = "3E6B8A"
BLUE_SOFT = "DCE7ED"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        edge_data = edges[edge]
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:" + key), str(edge_data[key]))


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ALUGUEL  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = rgb(INK_SOFT)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_document(doc: Document, preset: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6 if preset == "compact" else 8)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.333
    normal.paragraph_format.widow_control = True
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if preset == "compact" else WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, before, after in (
        ("Title", 26, 0, 10),
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14 if preset == "compact" else 12, 7 if preset == "compact" else 6),
        ("Heading 3", 12, 10 if preset == "compact" else 8, 5 if preset == "compact" else 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(COVER if style_name == "Title" else BRASS_DEEP)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        styles[list_name].font.name = "Calibri"
        styles[list_name].font.size = Pt(11)
        styles[list_name].paragraph_format.left_indent = Inches(0.375)
        styles[list_name].paragraph_format.first_line_indent = Inches(-0.188)
        styles[list_name].paragraph_format.space_after = Pt(4)
        styles[list_name].paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.208

    header = section.header.paragraphs[0]
    header.text = "ALUGUEL  /  GESTÃO DE ALUGUÉIS"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = rgb(BRASS_DEEP)
    header.paragraph_format.space_after = Pt(0)
    footer = section.footer.paragraphs[0]
    add_page_field(footer)


def add_text(paragraph, text: str, *, size=None, color=None, bold=False, italic=False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    run.font.bold = bold
    run.font.italic = italic


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    add_text(p, text.upper(), size=8.5, color=BRASS_DEEP, bold=True)


def add_lead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.25
    add_text(p, text, size=14, color=COVER, bold=True)


def add_rule(doc: Document, color=BRASS, width=18) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(width / 10)
    table.columns[1].width = Cm(15)
    set_cell_shading(table.cell(0, 0), color)
    set_cell_shading(table.cell(0, 1), PAPER)
    table.rows[0].height = Pt(3)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    remove_table_borders(table)


def add_chip_row(doc: Document, values: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(values))
    table.autofit = True
    for index, value in enumerate(values):
        cell = table.cell(0, index)
        set_cell_shading(cell, BRASS_SOFT)
        set_cell_margins(cell, 80, 100, 80, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, size=8.5, color=BRASS_DEEP, bold=True)
    remove_table_borders(table)


def add_callout(doc: Document, title: str, body: str, *, fill=BRASS_SOFT, accent=BRASS) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(0.18)
    table.columns[1].width = Cm(16.6)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    set_cell_margins(table.cell(0, 1), 150, 180, 150, 180)
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, size=10.5, color=COVER, bold=True)
    p2 = table.cell(0, 1).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, body, size=10, color=INK_SOFT)
    remove_table_borders(table)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_text(p, item)


def add_cover(doc: Document, kicker: str, title: str, subtitle: str, metadata: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17.2)
    cell.height = Cm(23.8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, COVER)
    set_cell_margins(cell, 550, 520, 500, 520)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_text(p, "⌂  ALUGUEL", size=13, color=BRASS, bold=True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_text(p, kicker.upper(), size=9, color=BRASS, bold=True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    add_text(p, title, size=29, color=CARD, bold=True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.line_spacing = 1.2
    add_text(p, subtitle, size=14, color="DCE8E2")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "—" * 11, size=10, color=BRASS, bold=True)
    for value in metadata:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_text(p, value, size=9, color="B9CCC3")
    remove_table_borders(table)
    doc.add_page_break()


def render_logo_assets() -> tuple[Path, Path]:
    BRANDING.mkdir(parents=True, exist_ok=True)
    dark = BRANDING / "logo-simbolo-dourado.png"
    light = BRANDING / "logo-simbolo-verde.png"
    scale = 32
    size = 32 * scale

    def make(path: Path, house: str, window: str, background=None) -> None:
        image = Image.new("RGBA", (size, size), background or (0, 0, 0, 0))
        draw = ImageDraw.Draw(image)
        def xy(points):
            return [(int(x * scale), int(y * scale)) for x, y in points]
        draw.polygon(xy([(16, 5), (27, 15), (5, 15)]), fill=house)
        draw.rounded_rectangle((8 * scale, 14 * scale, 24 * scale, 27 * scale), radius=scale, fill=house)
        draw.rectangle((21.5 * scale, 8 * scale, 25 * scale, 15 * scale), fill=house)
        draw.rectangle((11.5 * scale, 18.5 * scale, 15 * scale, 22 * scale), fill=window)
        draw.rectangle((17 * scale, 18.5 * scale, 20.5 * scale, 22 * scale), fill=window)
        image.save(path)

    make(dark, "#B8863C", "#14322A")
    make(light, "#14322A", "#F4F6F3")
    return dark, light


def add_logo_showcase(doc: Document, dark_logo: Path, light_logo: Path) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.2)
    table.columns[1].width = Cm(8.2)
    for cell, fill, logo, label, text_color in (
        (table.cell(0, 0), COVER, dark_logo, "Versão sobre fundo escuro", CARD),
        (table.cell(0, 1), PAPER, light_logo, "Versão sobre fundo claro", INK),
    ):
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 220, 220, 200, 220)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Cm(2.6))
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_text(p2, label, size=9, color=text_color, bold=True)
    remove_table_borders(table)


def add_palette(doc: Document) -> None:
    colors = [
        ("Verde capa", "#14322A", "Fundos, navegação, autoridade"),
        ("Verde apoio", "#1F4339", "Profundidade e superfícies escuras"),
        ("Latão", "#B8863C", "CTA, foco e acentos"),
        ("Papel", "#F4F6F3", "Fundo principal"),
        ("Tinta", "#1C2620", "Texto principal"),
        ("Ferrugem", "#A23B2E", "Atrasos e alertas críticos"),
        ("Manutenção", "#3E6B8A", "Status operacional"),
        ("Linha", "#DDE4DE", "Divisórias e bordas"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Cm(3.8)
    table.columns[1].width = Cm(3.2)
    table.columns[2].width = Cm(9.4)
    headers = ["Cor", "Código", "Uso recomendado"]
    for index, value in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, COVER)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, value, size=9, color=CARD, bold=True)
    set_repeat_table_header(table.rows[0])
    for name, code, usage in colors:
        cells = table.add_row().cells
        set_cell_shading(cells[0], code[1:])
        name_color = CARD if code in {"#14322A", "#1F4339", "#A23B2E", "#3E6B8A", "#1C2620"} else INK
        add_text(cells[0].paragraphs[0], name, size=9, color=name_color, bold=True)
        add_text(cells[1].paragraphs[0], code, size=9, color=INK, bold=True)
        add_text(cells[2].paragraphs[0], usage, size=9, color=INK_SOFT)
        for cell in cells:
            set_cell_margins(cell, 100, 120, 100, 120)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})


def build_brand_manual() -> Path:
    dark_logo, light_logo = render_logo_assets()
    doc = Document()
    configure_document(doc, "compact")
    add_cover(
        doc,
        "Mini manual de identidade comercial",
        "Aluguel",
        "Uma identidade clara para organizar, apresentar e vender a gestão de aluguéis.",
        ["Versão 1.0", CREATED, "Produto: Comercial 1.0"],
    )

    add_kicker(doc, "01  Essência da marca")
    doc.add_heading("Cada aluguel no lugar certo", level=1)
    add_lead(doc, "Clareza operacional para pequenos locadores que querem trocar planilhas e conversas dispersas por uma rotina centralizada.")
    doc.add_paragraph(
        "Aluguel é uma plataforma de gestão assistida que reúne imóveis, inquilinos, contratos, recebimentos, energia, despesas e interessados. A comunicação deve transmitir controle, proximidade e confiança — sem prometer automações ou resultados que o produto não executa."
    )
    add_callout(
        doc,
        "Proposta de valor principal",
        "Tenha cada imóvel, contrato, cobrança e documento no lugar certo — e saiba o que entrou, o que falta receber e o que exige atenção.",
    )
    doc.add_heading("Personalidade", level=2)
    add_chip_row(doc, ["Confiável", "Direta", "Organizada", "Próxima"])
    doc.add_heading("Posicionamento", level=2)
    doc.add_paragraph(
        "Para pequenos locadores que administram a própria carteira, Aluguel é a plataforma de gestão de locações que centraliza a operação da casa vaga ao pagamento recebido, com visão financeira, atalhos manuais de WhatsApp e PIX e portal do inquilino."
    )
    doc.add_page_break()

    add_kicker(doc, "02  Marca gráfica")
    doc.add_heading("Símbolo que nasce do produto", level=1)
    add_lead(doc, "A casa geométrica usada no aplicativo permanece como elemento principal. Nenhum novo símbolo foi inventado para os materiais comerciais.")
    add_logo_showcase(doc, dark_logo, light_logo)
    doc.add_heading("Área de proteção", level=2)
    doc.add_paragraph("Mantenha ao redor do símbolo uma área livre equivalente à largura de uma de suas janelas. Em lockups, preserve a relação de escala fornecida nos arquivos SVG.")
    doc.add_heading("Usos incorretos", level=2)
    add_bullets(doc, [
        "Não distorcer, inclinar ou alterar a proporção do símbolo.",
        "Não trocar as cores por efeitos, gradientes ou combinações fora da paleta.",
        "Não inserir detalhes, setas ou elementos que sugiram funcionalidades inexistentes.",
        "Não aplicar sobre fundos com pouco contraste.",
    ])
    doc.add_page_break()

    add_kicker(doc, "03  Paleta")
    doc.add_heading("Verde de controle, latão de ação", level=1)
    add_lead(doc, "A paleta vem diretamente do aplicativo: verde-escuro para confiança, tons de papel para leveza e latão para orientar a ação.")
    add_palette(doc)
    add_callout(doc, "Regra de contraste", "Textos pequenos devem usar Tinta sobre Papel ou branco sobre Verde capa. O latão funciona melhor como destaque, não como texto longo.", fill=BLUE_SOFT, accent=BLUE)
    doc.add_page_break()

    add_kicker(doc, "04  Tipografia")
    doc.add_heading("Tecnologia com leitura humana", level=1)
    add_lead(doc, "A combinação tipográfica do produto deve ser preservada na comunicação digital.")
    typography = doc.add_table(rows=3, cols=2)
    typography.autofit = False
    typography.columns[0].width = Cm(5.2)
    typography.columns[1].width = Cm(11.2)
    rows = [
        ("Space Grotesk", "Títulos, chamadas, números de impacto e nome Aluguel."),
        ("IBM Plex Sans", "Textos, legendas, formulários, FAQs e documentos digitais."),
        ("IBM Plex Mono", "Valores, datas, chips, versões e informações operacionais."),
    ]
    for index, (font, usage) in enumerate(rows):
        c1, c2 = typography.rows[index].cells
        set_cell_shading(c1, COVER if index == 0 else PAPER)
        set_cell_shading(c2, BRASS_SOFT if index == 0 else CARD)
        add_text(c1.paragraphs[0], font, size=13 if index == 0 else 11, color=CARD if index == 0 else COVER, bold=True)
        add_text(c2.paragraphs[0], usage, size=10, color=INK_SOFT)
        for cell in (c1, c2):
            set_cell_margins(cell, 180, 180, 180, 180)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_heading("Fallbacks", level=2)
    doc.add_paragraph("Quando as fontes web não estiverem disponíveis, use Calibri ou Arial em documentos de escritório. Evite fontes decorativas, manuscritas ou excessivamente condensadas.")
    doc.add_page_break()

    add_kicker(doc, "05  Voz e mensagem")
    doc.add_heading("Fale como quem organiza a rotina", level=1)
    add_lead(doc, "Frases curtas, benefícios concretos e transparência sobre o que depende de ação humana.")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.2)
    table.columns[1].width = Cm(8.2)
    for index, (title, fill, accent) in enumerate((("Preferir", "EAF5ED", "21613C"), ("Evitar", RUST_SOFT, RUST))):
        cell = table.cell(0, index)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 180, 190, 180, 190)
        add_text(cell.paragraphs[0], title, size=12, color=accent, bold=True)
        items = ([
            "“Veja o que falta receber.”",
            "“Cobrança manual com mensagem pronta.”",
            "“Dados separados por conta.”",
            "“Portal de consulta para o inquilino.”",
        ] if index == 0 else [
            "“Acabe com a inadimplência.”",
            "“Cobrança 100% automática.”",
            "“Conciliação bancária.”",
            "“Garantia de conformidade LGPD.”",
        ])
        for item in items:
            p = cell.add_paragraph(style="List Bullet")
            add_text(p, item, size=9.5, color=INK)
    remove_table_borders(table)
    doc.add_heading("Hierarquia de mensagens", level=2)
    add_numbered(doc, [
        "Controle: seus aluguéis em uma visão única.",
        "Rotina: imóveis, contratos, cobranças e documentos conectados.",
        "Decisão: previsto, recebido, pendências, ocupação e despesas.",
        "Relacionamento: atalhos manuais para WhatsApp/PIX e portal do inquilino.",
    ])
    doc.add_page_break()

    add_kicker(doc, "06  Sistema visual")
    doc.add_heading("Do produto para a comunicação", level=1)
    add_lead(doc, "Materiais comerciais devem parecer uma extensão natural do aplicativo.")
    add_bullets(doc, [
        "Fundos em Papel com blocos Verde capa; cartões brancos com borda Linha.",
        "Cantos entre 16 e 22 px no digital; sombras suaves e pouco contrastadas.",
        "Chips arredondados para status e categorias; números em IBM Plex Mono.",
        "Fotografias e mockups sempre acompanhados de telas reais ou demonstrativas renderizadas pelo produto.",
        "Ferrugem apenas para atraso/risco; azul apenas para manutenção ou contexto operacional.",
    ])
    add_callout(doc, "Direção de arte", "Premium não significa ornamental. O sistema visual deve priorizar leitura, hierarquia, espaço e prova real do produto.", fill=PAPER, accent=BRASS)
    doc.add_heading("Composição recomendada", level=2)
    composition = doc.add_table(rows=1, cols=3)
    composition.autofit = False
    composition.columns[0].width = Cm(5.45)
    composition.columns[1].width = Cm(5.45)
    composition.columns[2].width = Cm(5.45)
    for cell, title, body, fill in zip(composition.rows[0].cells, ["1  Contexto", "2  Produto", "3  Ação"], ["Nomeie a dor sem exagero.", "Mostre a tela que resolve.", "Finalize com CTA único."], [PAPER, BRASS_SOFT, COVER]):
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 190, 170, 190, 170)
        text_color = CARD if fill == COVER else COVER
        add_text(cell.paragraphs[0], title, size=11, color=text_color, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_text(p, body, size=9, color="DCE8E2" if fill == COVER else INK_SOFT)
    remove_table_borders(composition)
    doc.add_page_break()

    add_kicker(doc, "07  Chamadas e CTAs")
    doc.add_heading("Uma ação principal por peça", level=1)
    add_lead(doc, "A demonstração é o próximo passo comercial principal; criar a conta gratuita é a alternativa de autosserviço.")
    cta = doc.add_table(rows=3, cols=2)
    cta.autofit = False
    cta.columns[0].width = Cm(5.2)
    cta.columns[1].width = Cm(11.2)
    for row, (label, copy, fill) in zip(cta.rows, [
        ("CTA principal", "Solicitar uma demonstração", COVER),
        ("CTA secundário", "Começar grátis", BRASS_SOFT),
        ("Microcopy", "Demonstração guiada. Sem compromisso. Condições sob consulta.", PAPER),
    ]):
        set_cell_shading(row.cells[0], fill)
        set_cell_shading(row.cells[1], fill)
        text_color = CARD if fill == COVER else COVER
        add_text(row.cells[0].paragraphs[0], label, size=9, color=text_color, bold=True)
        add_text(row.cells[1].paragraphs[0], copy, size=10.5, color=text_color, bold=(label != "Microcopy"))
        for cell in row.cells:
            set_cell_margins(cell, 140, 160, 140, 160)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_heading("Destino oficial", level=2)
    doc.add_paragraph(APP_URL)
    add_callout(doc, "WhatsApp pendente", "O botão usa e-mail como fallback funcional até a confirmação do número comercial com DDD. Nenhum número fictício deve ser publicado.", fill=RUST_SOFT, accent=RUST)
    doc.add_page_break()

    add_kicker(doc, "08  Governança")
    doc.add_heading("Consistência antes de velocidade", level=1)
    add_lead(doc, "Toda peça deve passar por quatro verificações antes de publicação.")
    checks = [
        ("Verdade", "A funcionalidade aparece no produto e foi validada?"),
        ("Privacidade", "A tela usa dados demonstrativos e não expõe pessoas reais?"),
        ("Legibilidade", "Texto, CTA, contraste e QR Code funcionam no tamanho final?"),
        ("Destino", "Links, e-mail, WhatsApp e URL estão corretos?"),
    ]
    for title, body in checks:
        add_callout(doc, title, body, fill=PAPER, accent=BRASS)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)
    doc.add_heading("Arquivos-base", level=2)
    add_bullets(doc, [
        "SVGs do símbolo e lockups: versões editáveis para web e design.",
        "brand-tokens.css / brand-tokens.json: referência técnica da identidade.",
        "Este manual em DOCX e PDF: orientação de uso e governança.",
        "README.md: índice dos arquivos e data de criação.",
    ])
    add_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, f"Aluguel  •  {EMAIL}  •  {CREATED}", size=9, color=INK_SOFT, bold=True)

    output = BRANDING / "manual-identidade-comercial-aluguel.docx"
    doc.save(output)
    return output


def add_section_start(doc: Document, number: str, title: str, lead: str) -> None:
    add_kicker(doc, number)
    doc.add_heading(title, level=1)
    add_lead(doc, lead)


def add_metadata_grid(doc: Document, left: list[tuple[str, str]], right: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=max(len(left), len(right)), cols=4)
    table.autofit = False
    widths = [Cm(2.6), Cm(5.3), Cm(2.6), Cm(5.9)]
    for index, width in enumerate(widths):
        table.columns[index].width = width
    for row_index in range(len(table.rows)):
        pairs = [left[row_index] if row_index < len(left) else ("", ""), right[row_index] if row_index < len(right) else ("", "")]
        for pair_index, pair in enumerate(pairs):
            label_cell = table.cell(row_index, pair_index * 2)
            value_cell = table.cell(row_index, pair_index * 2 + 1)
            set_cell_shading(label_cell, PAPER)
            set_cell_shading(value_cell, CARD)
            add_text(label_cell.paragraphs[0], pair[0], size=8.5, color=INK_SOFT, bold=True)
            add_text(value_cell.paragraphs[0], pair[1], size=9, color=INK)
            for cell in (label_cell, value_cell):
                set_cell_margins(cell, 90, 100, 90, 100)
                set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})


def build_proposal() -> Path:
    PROPOSAL.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc, "narrative")
    add_cover(
        doc,
        "Proposta comercial",
        "Aluguel",
        "Gestão de imóveis, contratos e recebimentos para uma rotina mais organizada.",
        ["Preparada para: [NOME DO CLIENTE]", "Validade: [INSERIR VALIDADE]", CREATED],
    )

    add_kicker(doc, "Proposta para [NOME DO CLIENTE]")
    doc.add_heading("Uma operação de locação mais clara", level=1)
    add_lead(doc, "Esta proposta apresenta como o Aluguel pode centralizar a gestão cotidiana da carteira, do imóvel vago ao acompanhamento do recebimento.")
    add_metadata_grid(doc, [
        ("Cliente", "[NOME / EMPRESA]"),
        ("Responsável", "[NOME DO DECISOR]"),
        ("Carteira", "[NÚMERO DE IMÓVEIS]"),
    ], [
        ("Proponente", "Aluguel"),
        ("Contato", EMAIL),
        ("Proposta", "[CÓDIGO / VERSÃO]"),
    ])
    doc.add_heading("Resumo executivo", level=2)
    doc.add_paragraph(
        "O Aluguel reúne imóveis, inquilinos, contratos, recebimentos, energia, despesas, interessados, documentos e agenda em uma única plataforma. O objetivo é oferecer visão operacional e financeira, preservar histórico e facilitar ações manuais de cobrança por WhatsApp e PIX — sem substituir análise jurídica, contábil ou conciliação bancária."
    )
    add_callout(doc, "Resultado esperado", "Uma rotina mais organizada, com informações localizáveis e pendências visíveis. Indicadores de resultado serão definidos e medidos junto ao cliente; não há percentuais garantidos.")
    doc.add_page_break()

    add_section_start(doc, "01  Contexto", "O cenário que esta solução endereça", "Quando imóveis, cobranças e documentos ficam espalhados, a dificuldade não é apenas registrar: é saber rapidamente o que exige atenção.")
    doc.add_heading("Situação típica", level=2)
    add_bullets(doc, [
        "Cadastros de imóveis e inquilinos distribuídos entre planilhas, anotações e conversas.",
        "Pagamentos, energia e despesas sem uma visão consolidada por mês e por imóvel.",
        "Histórico contratual difícil de consultar quando há troca de inquilino.",
        "Interessados e imóveis vagos acompanhados sem um funil único.",
        "Documentos e fotos sem vínculo claro com a operação correspondente.",
    ])
    doc.add_heading("Ação proposta", level=2)
    doc.add_paragraph("Implantar uma conta Aluguel compatível com o tamanho da carteira, configurar os dados do locador, organizar os cadastros prioritários e orientar o responsável pela operação nos fluxos essenciais.")
    doc.add_heading("Resultado operacional buscado", level=2)
    doc.add_paragraph("Concentrar a rotina em um ambiente com status, históricos, alertas e relatórios gerenciais. O registro e a confirmação de recebimentos permanecem sob responsabilidade do usuário.")
    doc.add_page_break()

    add_section_start(doc, "02  Solução", "O que o Aluguel entrega", "Uma plataforma web instalável, desenhada para a gestão assistida de pequenos locadores e carteiras enxutas.")
    features = [
        ("Painel", "Previsto, recebido, falta receber, ocupação, alertas e visão de 12 meses."),
        ("Imóveis", "Cadastro, status, características, contratos, pagamentos, energia, despesas, fotos e documentos."),
        ("Inquilinos", "Base reutilizável, vínculos e histórico contratual por ciclo de ocupação."),
        ("Financeiro", "Acompanhamento mensal/anual, ageing, despesas, saldo e exportações CSV/PDF."),
        ("Relacionamento", "Mensagens pré-preenchidas para WhatsApp, PIX Copia e Cola e portal de consulta do inquilino."),
        ("Vagas", "Catálogo público de imóveis vagos/publicados e funil de interessados com combinação por regras."),
        ("Operação", "Calendário, lembretes manuais, equipe, backup exportável e consulta offline da última cópia."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.3)
    table.columns[1].width = Cm(12.1)
    for index, title in enumerate(("Módulo", "Capacidade")):
        set_cell_shading(table.cell(0, index), COVER)
        add_text(table.cell(0, index).paragraphs[0], title, size=9, color=CARD, bold=True)
    set_repeat_table_header(table.rows[0])
    for name, desc in features:
        cells = table.add_row().cells
        set_cell_shading(cells[0], BRASS_SOFT)
        add_text(cells[0].paragraphs[0], name, size=9.5, color=COVER, bold=True)
        add_text(cells[1].paragraphs[0], desc, size=9.5, color=INK_SOFT)
        for cell in cells:
            set_cell_margins(cell, 100, 120, 100, 120)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_page_break()

    add_section_start(doc, "03  Escopo", "Configuração e ativação", "O escopo comercial deve ser confirmado conforme o volume de dados e o nível de apoio desejado.")
    doc.add_heading("Incluído nesta proposta", level=2)
    add_bullets(doc, [
        "Licença de uso da plataforma no plano selecionado.",
        "Criação da conta e configuração inicial do perfil do locador.",
        "Orientação de início para cadastro de imóveis, inquilinos e contratos.",
        "Acesso aos módulos disponíveis no plano e às atualizações da versão contratada.",
        "[INSERIR ESCOPO DE TREINAMENTO / MIGRAÇÃO / SUPORTE].",
    ])
    doc.add_heading("Não incluído", level=2)
    add_bullets(doc, [
        "Serviço jurídico, contábil, fiscal ou elaboração/assinatura de contratos.",
        "Gateway de pagamento, conciliação bancária ou confirmação automática de PIX.",
        "Envio automático de WhatsApp, notificações ou régua de cobrança.",
        "Integrações personalizadas ou migração de dados não descritas no escopo.",
        "Garantia de redução de inadimplência ou de resultado financeiro específico.",
    ])
    add_callout(doc, "Responsabilidade compartilhada", "O cliente valida os cadastros, registra recebimentos, mantém dados de contato atualizados e define quem pode acessar a conta.", fill=BLUE_SOFT, accent=BLUE)
    doc.add_page_break()

    add_section_start(doc, "04  Planos", "Escolha compatível com a carteira", "Os limites abaixo vêm da configuração atual do produto. Valores, periodicidade e política comercial precisam ser preenchidos antes do envio.")
    plans = doc.add_table(rows=1, cols=4)
    plans.autofit = False
    widths = [Cm(3.1), Cm(3.7), Cm(3.5), Cm(6.1)]
    for index, width in enumerate(widths):
        plans.columns[index].width = width
    for index, title in enumerate(("Plano", "Imóveis", "Armazenamento", "Investimento")):
        set_cell_shading(plans.cell(0, index), COVER)
        add_text(plans.cell(0, index).paragraphs[0], title, size=9, color=CARD, bold=True)
    for name, houses, storage, price in (
        ("Gratuito", "1", "50 MB", "R$ 0,00 / conforme política vigente"),
        ("Básico", "3", "1 GB", "[INSERIR VALOR E PERIODICIDADE]"),
        ("Premium", "100", "10 GB", "[INSERIR VALOR E PERIODICIDADE]"),
    ):
        cells = plans.add_row().cells
        set_cell_shading(cells[0], BRASS_SOFT if name == "Básico" else PAPER)
        add_text(cells[0].paragraphs[0], name, size=10, color=COVER, bold=True)
        add_text(cells[1].paragraphs[0], houses, size=10, color=INK)
        add_text(cells[2].paragraphs[0], storage, size=10, color=INK)
        add_text(cells[3].paragraphs[0], price, size=9, color=INK_SOFT, bold=(name == "Básico"))
        for cell in cells:
            set_cell_margins(cell, 120, 120, 120, 120)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_heading("Plano recomendado", level=2)
    doc.add_paragraph("[PLANO] — recomendado com base em [NÚMERO DE IMÓVEIS], [VOLUME DE ARQUIVOS] e [NÚMERO DE USUÁRIOS].")
    doc.add_heading("Condições comerciais", level=2)
    add_metadata_grid(doc, [
        ("Investimento", "[R$ / PERÍODO]"),
        ("Pagamento", "[FORMA E VENCIMENTO]"),
        ("Reajuste", "[ÍNDICE / REGRA]"),
    ], [
        ("Ativação", "[PRAZO]"),
        ("Cancelamento", "[POLÍTICA]"),
        ("Validade", "[DATA / DIAS]"),
    ])
    doc.add_page_break()

    add_section_start(doc, "05  Implantação", "Próximos passos claros", "A ativação final depende do escopo de apoio e das informações que o cliente fornecer.")
    add_numbered(doc, [
        "Aprovação da proposta e confirmação das condições comerciais.",
        "Criação ou validação da conta do responsável.",
        "Configuração de perfil, PIX e preferências operacionais.",
        "Cadastro ou importação assistida do escopo acordado.",
        "Orientação dos fluxos essenciais e validação de acesso.",
        "Início da operação e acompanhamento conforme plano de suporte.",
    ])
    doc.add_heading("Cronograma", level=2)
    timeline = doc.add_table(rows=1, cols=3)
    timeline.autofit = False
    timeline.columns[0].width = Cm(4.5)
    timeline.columns[1].width = Cm(7.2)
    timeline.columns[2].width = Cm(4.7)
    for index, title in enumerate(("Etapa", "Entrega", "Prazo")):
        set_cell_shading(timeline.cell(0, index), COVER)
        add_text(timeline.cell(0, index).paragraphs[0], title, size=9, color=CARD, bold=True)
    for row_data in (
        ("Preparação", "Conta e configurações", "[PRAZO]"),
        ("Organização", "Cadastros do escopo", "[PRAZO]"),
        ("Orientação", "Sessão de início", "[PRAZO]"),
        ("Operação", "Acompanhamento acordado", "[PRAZO]"),
    ):
        cells = timeline.add_row().cells
        for index, value in enumerate(row_data):
            add_text(cells[index].paragraphs[0], value, size=9.5, color=INK if index != 1 else INK_SOFT, bold=(index == 0))
            set_cell_margins(cells[index], 100, 120, 100, 120)
            set_cell_border(cells[index], bottom={"val": "single", "sz": "4", "color": LINE})
    add_callout(doc, "Pendência a preencher", "Defina prazo de ativação, quantidade de encontros, canal/horário de suporte e responsabilidades de migração antes de enviar esta proposta.", fill=RUST_SOFT, accent=RUST)
    doc.add_page_break()

    add_section_start(doc, "06  Confiança", "Como os dados são tratados no produto", "O aplicativo utiliza separação por conta e políticas de acesso no banco; a comunicação deve permanecer precisa e sem alegar certificações não verificadas.")
    add_bullets(doc, [
        "Dados operacionais separados por proprietário e perfis de acesso.",
        "Arquivos privados entregues por links temporários quando autorizados.",
        "Backup manual exportável e snapshots diários conforme a configuração atual.",
        "Portal do inquilino em modo de consulta, limitado aos dados e documentos liberados.",
        "Consulta offline da última cópia disponível no aparelho; alterações exigem internet.",
    ])
    doc.add_heading("Observações", level=2)
    doc.add_paragraph("A solução não deve ser apresentada como certificada em LGPD, inviolável ou disponível sem interrupção. Termos de uso, aviso de privacidade e responsabilidades contratuais devem passar por revisão jurídica antes da comercialização em escala.")
    doc.add_heading("Suporte", level=2)
    doc.add_paragraph("Canal: [E-MAIL / WHATSAPP]  •  Horário: [JANELA]  •  Prazo inicial de resposta: [SLA]  •  Escopo: [DESCREVER].")
    doc.add_page_break()

    add_section_start(doc, "07  Aceite", "Vamos organizar sua carteira?", "Após o preenchimento dos campos comerciais, este documento estará pronto para assinatura e início do processo de ativação.")
    add_callout(doc, "Ação recomendada", "Agendar uma demonstração guiada e validar o plano adequado ao número de imóveis e ao volume de arquivos.")
    doc.add_heading("Aceite da proposta", level=2)
    accept = doc.add_table(rows=4, cols=2)
    accept.autofit = False
    accept.columns[0].width = Cm(8.2)
    accept.columns[1].width = Cm(8.2)
    labels = [
        ("Cliente / Razão social", "CPF / CNPJ"),
        ("Nome do responsável", "Cargo / função"),
        ("Assinatura", "Data"),
        ("Plano selecionado", "Condição comercial"),
    ]
    for row, pair in zip(accept.rows, labels):
        for cell, label in zip(row.cells, pair):
            set_cell_shading(cell, PAPER)
            set_cell_margins(cell, 110, 140, 330, 140)
            add_text(cell.paragraphs[0], label, size=8.5, color=INK_SOFT, bold=True)
            set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": LINE}, left={"val": "single", "sz": "4", "color": LINE}, right={"val": "single", "sz": "4", "color": LINE}, top={"val": "single", "sz": "4", "color": LINE})
    doc.add_heading("Contato", level=2)
    doc.add_paragraph(f"E-mail: {EMAIL}\nAplicativo: {APP_URL}\nWhatsApp: [INSERIR NÚMERO COM DDD]")
    add_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Aluguel — cada aluguel no lugar certo.", size=11, color=COVER, bold=True)

    output = PROPOSAL / "proposta-comercial-aluguel-editavel.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    brand = build_brand_manual()
    proposal = build_proposal()
    print(brand)
    print(proposal)
